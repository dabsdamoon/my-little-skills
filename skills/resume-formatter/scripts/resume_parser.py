#!/usr/bin/env python3
"""
Resume Parser - Robust Content Extraction

Extracts structured content from diverse resume formats including:
- Multi-column layouts
- Table-based resumes
- Headers/footers/text boxes
- Korean and English resumes
- Mixed formats

Usage:
    python resume_parser.py <resume_file.docx> [--output parsed.json]
"""

import argparse
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    exit(1)


class ResumeParser:
    """Robust resume content extractor"""

    # Section heading synonyms
    SECTION_SYNONYMS = {
        'experience': [
            'experience', 'work experience', 'professional experience',
            'employment', 'work history', 'career history', 'professional background',
            '경력', '경력사항', '상세경력사항', '직장경력', '근무경력'
        ],
        'education': [
            'education', 'academic background', 'academic history',
            'educational background', 'qualifications',
            '학력', '학력사항', '교육사항'
        ],
        'skills': [
            'skills', 'technical skills', 'core competencies', 'expertise',
            'proficiencies', 'competencies', 'technical proficiencies',
            '핵심역량', '보유기술', '전문기술', '기술스택'
        ],
        'summary': [
            'summary', 'professional summary', 'profile', 'objective',
            'career objective', 'about me',
            '자기소개', '요약'
        ],
        'personal': [
            'personal information', 'contact', 'contact information',
            '인적사항', '개인정보'
        ]
    }

    def __init__(self, resume_path: str):
        self.resume_path = Path(resume_path)
        self.doc = Document(str(self.resume_path))
        self.extracted = {
            'personal_info': {},
            'experience': [],
            'education': [],
            'skills': {},
            'summary': None,
            'metadata': {
                'source_file': self.resume_path.name,
                'extraction_date': datetime.now().isoformat(),
                'warnings': []
            }
        }

    def parse(self) -> Dict[str, Any]:
        """Main parsing workflow"""
        print(f"Parsing resume: {self.resume_path.name}")

        # Multi-pass extraction
        self._extract_from_headers_footers()
        self._extract_from_body()
        self._extract_from_tables()
        self._extract_from_textboxes()

        # Post-processing
        self._deduplicate()
        self._calculate_confidence()

        return self.extracted

    def _extract_from_headers_footers(self):
        """Extract from headers and footers (often contains contact info)"""
        print("  Checking headers and footers...")

        for section in self.doc.sections:
            # Header
            for para in section.header.paragraphs:
                text = para.text.strip()
                if text:
                    self._extract_contact_info(text, source='header')

            # Footer
            for para in section.footer.paragraphs:
                text = para.text.strip()
                if text:
                    self._extract_contact_info(text, source='footer')
                    # Portfolio links often in footer
                    urls = re.findall(r'https?://[^\s]+', text)
                    for url in urls:
                        if 'linkedin' in url.lower():
                            self.extracted['personal_info']['linkedin'] = {
                                'value': url,
                                'confidence': 0.9,
                                'source': 'footer'
                            }
                        elif 'github' in url.lower():
                            self.extracted['personal_info']['github'] = {
                                'value': url,
                                'confidence': 0.9,
                                'source': 'footer'
                            }

    def _extract_from_body(self):
        """Extract from main document body"""
        print("  Parsing main body...")

        current_section = None
        buffer = []

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a section heading
            section = self._detect_section(text, para)
            if section:
                # Process buffered content from previous section
                if current_section and buffer:
                    self._process_section_content(current_section, buffer)
                    buffer = []
                current_section = section
                continue

            # Add to current section buffer
            if current_section:
                buffer.append(text)
            else:
                # No section yet - might be name/contact at top
                self._extract_contact_info(text, source='body')

        # Process final section
        if current_section and buffer:
            self._process_section_content(current_section, buffer)

    def _extract_from_tables(self):
        """Extract from table-based resume structures"""
        if not self.doc.tables:
            return

        print(f"  Parsing {len(self.doc.tables)} tables...")

        for table in self.doc.tables:
            self._process_table(table)

    def _extract_from_textboxes(self):
        """Extract from text boxes and shapes"""
        print("  Checking for text boxes...")

        # Note: Accessing text boxes requires XML parsing
        # This is a simplified version - full implementation would parse document.xml
        try:
            for element in self.doc.element.body.iter():
                # Look for text frames in shapes
                if element.tag.endswith('txBody'):
                    for text_elem in element.iter():
                        if text_elem.text:
                            # Process text box content
                            self._extract_contact_info(text_elem.text, source='textbox')
        except Exception as e:
            self.extracted['metadata']['warnings'].append(
                f"Could not extract text boxes: {str(e)}"
            )

    def _detect_section(self, text: str, para: Paragraph) -> Optional[str]:
        """Detect if paragraph is a section heading"""
        text_lower = text.lower()

        # Check formatting (headings are often bold, larger, or styled)
        is_formatted = False
        if para.runs:
            # Check if bold
            is_formatted = any(run.bold for run in para.runs)

        # Check against synonyms
        for section_type, synonyms in self.SECTION_SYNONYMS.items():
            for synonym in synonyms:
                if text_lower == synonym or text_lower.startswith(synonym + ':'):
                    return section_type
                # Partial match if formatted like a heading
                if is_formatted and synonym in text_lower:
                    return section_type

        return None

    def _process_section_content(self, section: str, content: List[str]):
        """Process content for a specific section"""
        if section == 'experience':
            self._parse_experience(content)
        elif section == 'education':
            self._parse_education(content)
        elif section == 'skills':
            self._parse_skills(content)
        elif section == 'summary':
            self.extracted['summary'] = ' '.join(content)

    def _parse_experience(self, lines: List[str]):
        """Parse work experience section with project detection

        Handles two formats:
        1. Standard: "Company | Title | Dates" all on one line
        2. Korean format: Date on one line, then "Company - Title" on next line
        """
        current_job = None
        current_project = None
        pending_dates = None  # Store dates found on separate line
        i = 0

        while i < len(lines):
            line = lines[i]

            # Check if this line is a standalone date (Korean format)
            if self._is_standalone_date_line(line) and i + 1 < len(lines):
                # Save dates and check next line for company
                pending_dates = self._parse_date_range(line)
                next_line = lines[i + 1]

                # Check if next line is company info
                if self._looks_like_company_line(next_line):
                    # Save current job if exists
                    if current_job:
                        if current_project:
                            if 'projects' not in current_job:
                                current_job['projects'] = []
                            current_job['projects'].append(current_project)
                            current_project = None
                        self.extracted['experience'].append(current_job)

                    # Create new job with dates from previous line
                    current_job = self._extract_company_title(next_line)
                    current_job.update(pending_dates)
                    current_job['confidence'] = 0.9  # High confidence for Korean format
                    pending_dates = None
                    i += 2  # Skip both lines
                    continue

            # Try to detect job entry (single-line format)
            elif self._looks_like_job_header(line):
                if current_job:
                    if current_project:
                        if 'projects' not in current_job:
                            current_job['projects'] = []
                        current_job['projects'].append(current_project)
                        current_project = None
                    self.extracted['experience'].append(current_job)
                current_job = self._extract_job_info(line)

            elif current_job:
                # Check if this is a project header within the job
                if self._looks_like_project_header(line):
                    if current_project:
                        # Save previous project
                        if 'projects' not in current_job:
                            current_job['projects'] = []
                        current_job['projects'].append(current_project)
                    current_project = {
                        'name': line.strip('[] 【】()'),
                        'responsibilities': []
                    }
                elif current_project:
                    # Add to current project responsibilities
                    clean_line = re.sub(r'^[•\-\*]\s*', '', line)
                    if clean_line:
                        current_project['responsibilities'].append(clean_line)
                else:
                    # Add to general job responsibilities
                    if 'responsibilities' not in current_job:
                        current_job['responsibilities'] = []
                    clean_line = re.sub(r'^[•\-\*]\s*', '', line)
                    if clean_line:
                        current_job['responsibilities'].append(clean_line)

            i += 1

        # Add final project if exists
        if current_project and current_job:
            if 'projects' not in current_job:
                current_job['projects'] = []
            current_job['projects'].append(current_project)

        # Add final job
        if current_job:
            self.extracted['experience'].append(current_job)

    def _is_standalone_date_line(self, line: str) -> bool:
        """Check if line contains only date range (Korean format)"""
        # Korean format: "2023년 7월 - 현재" or "2023년7월 - 현재" (space optional)
        korean_date = re.match(r'^\s*\d{4}년\s*\d{1,2}월?\s*[-~]\s*(현재|재직중|Present|\d{4}년\s*\d{1,2}월?)\s*$', line)
        if korean_date:
            return True

        # Dot format: "2023.07 - 현재"
        dot_date = re.match(r'^\s*\d{4}\.\d{2}\s*[-~]\s*(현재|재직중|Present|\d{4}\.\d{2})\s*$', line)
        return bool(dot_date)

    def _looks_like_company_line(self, line: str) -> bool:
        """Check if line looks like a company name with title"""
        # Contains company indicators
        has_company_indicator = bool(re.search(
            r'(Inc\.|Ltd\.|Corp\.|Co\.|주식회사|Company|AI|엔지니어|연구원|개발자|Part)',
            line,
            re.I
        ))

        # Contains dash separator (Company - Title format)
        has_dash_separator = ' - ' in line or ',' in line

        # Not too long (company+title lines are usually concise)
        not_too_long = len(line) < 100

        return (has_company_indicator or has_dash_separator) and not_too_long

    def _extract_company_title(self, line: str) -> Dict[str, Any]:
        """Extract company and title from Korean format line"""
        job = {'confidence': 0.85}

        # Try dash separator: "Company - Title"
        if ' - ' in line:
            parts = line.split(' - ', 1)
            job['company'] = parts[0].strip()
            if len(parts) > 1:
                job['title'] = parts[1].strip()
        # Try comma separator: "Company, Department - Title"
        elif ',' in line:
            parts = line.split(',')
            job['company'] = parts[0].strip()
            if len(parts) > 1:
                # Rest might be department and title
                rest = ','.join(parts[1:]).strip()
                if ' - ' in rest:
                    dept_title = rest.split(' - ')
                    job['department'] = dept_title[0].strip()
                    if len(dept_title) > 1:
                        job['title'] = dept_title[1].strip()
                else:
                    job['title'] = rest
        else:
            # Fallback: entire line is company
            job['company'] = line.strip()

        return job

    def _looks_like_project_header(self, line: str) -> bool:
        """Detect if line is a project header within job experience"""
        # Projects often marked with brackets or special formatting
        if line.startswith('[') or line.startswith('【'):
            return True

        # Common project keywords
        project_keywords = ['프로젝트', 'project', '과제', 'task force']
        if any(kw in line.lower() for kw in project_keywords):
            # But not if it's in a sentence
            if len(line.split()) <= 6:  # Short line more likely to be header
                return True

        # Date ranges within experience might indicate project
        if re.search(r'\d{4}\.\d{2}\s*[-~]\s*\d{4}\.\d{2}', line) and len(line) < 100:
            return True

        return False

    def _looks_like_job_header(self, line: str) -> bool:
        """Heuristic to detect job entry start"""
        # Contains dates?
        has_dates = bool(re.search(r'\d{4}', line))
        # Contains company indicators?
        has_company = bool(re.search(r'(Inc\.|Ltd\.|Corp\.|Co\.|주식회사|Company)', line, re.I))
        # Contains common title words?
        has_title = bool(re.search(r'(Engineer|Developer|Manager|Analyst|Researcher|팀장|연구원)', line, re.I))

        return (has_dates and (has_company or has_title)) or (line.count('|') >= 2)

    def _extract_job_info(self, line: str) -> Dict[str, Any]:
        """Extract company, title, dates from job header"""
        job = {'confidence': 0.7}

        # Try pipe-separated format: Company | Title | Dates
        if '|' in line:
            parts = [p.strip() for p in line.split('|')]
            if len(parts) >= 2:
                job['company'] = parts[0]
                job['title'] = parts[1]
                if len(parts) >= 3:
                    dates = self._parse_date_range(parts[2])
                    job.update(dates)
                job['confidence'] = 0.85
                return job

        # Try to extract dates
        dates = self._parse_date_range(line)
        job.update(dates)

        # Remaining text might be company/title
        text_without_dates = re.sub(r'\d{4}[년\./-]\s*\d{1,2}[월]?\s*[-~]\s*(현재|Present|\d{4})', '', line)
        text_without_dates = re.sub(r'\d{4}\s*-\s*\d{4}', '', text_without_dates)

        # Simple heuristic: first part is company, second is title
        parts = [p.strip() for p in text_without_dates.split('-') if p.strip()]
        if len(parts) >= 2:
            job['company'] = parts[0]
            job['title'] = parts[1]
        elif parts:
            job['company'] = parts[0]

        return job

    def _parse_date_range(self, text: str) -> Dict[str, str]:
        """Extract start and end dates from text"""
        dates = {}

        # Korean format: YYYY년 MM월
        korean_match = re.search(r'(\d{4})년\s*(\d{1,2})월?\s*[-~]\s*(현재|재직중|Present|\d{4}년\s*\d{1,2}월?)', text)
        if korean_match:
            dates['start_date'] = f"{korean_match.group(1)}.{korean_match.group(2):0>2}"
            end = korean_match.group(3)
            if end in ['현재', '재직중', 'Present']:
                dates['end_date'] = 'Present'
            else:
                end_match = re.search(r'(\d{4})년\s*(\d{1,2})월?', end)
                if end_match:
                    dates['end_date'] = f"{end_match.group(1)}.{end_match.group(2):0>2}"
            return dates

        # Dot format: YYYY.MM
        dot_match = re.search(r'(\d{4})\.(\d{2})\s*[-~]\s*(현재|재직중|Present|\d{4}\.\d{2})', text)
        if dot_match:
            dates['start_date'] = f"{dot_match.group(1)}.{dot_match.group(2)}"
            end = dot_match.group(3)
            dates['end_date'] = 'Present' if end in ['현재', '재직중', 'Present'] else end
            return dates

        # Western format: Month YYYY or MM/YYYY
        western_match = re.search(r'(\w+)\s+(\d{4})\s*[-–]\s*(Present|Current|\w+\s+\d{4})', text)
        if western_match:
            dates['start_date'] = f"{western_match.group(2)}.{self._month_to_num(western_match.group(1))}"
            end = western_match.group(3)
            if end in ['Present', 'Current']:
                dates['end_date'] = 'Present'
            else:
                end_match = re.search(r'(\w+)\s+(\d{4})', end)
                if end_match:
                    dates['end_date'] = f"{end_match.group(2)}.{self._month_to_num(end_match.group(1))}"
            return dates

        return dates

    def _month_to_num(self, month: str) -> str:
        """Convert month name to number"""
        months = {
            'jan': '01', 'feb': '02', 'mar': '03', 'apr': '04',
            'may': '05', 'jun': '06', 'jul': '07', 'aug': '08',
            'sep': '09', 'oct': '10', 'nov': '11', 'dec': '12'
        }
        return months.get(month.lower()[:3], '01')

    def _parse_education(self, lines: List[str]):
        """Parse education section"""
        for line in lines:
            # Look for university/school names
            if any(keyword in line.lower() for keyword in ['university', 'college', 'institute', '대학교', '대학']):
                edu = {'confidence': 0.75}

                # Extract GPA
                gpa_match = re.search(r'(\d\.\d+)[/\s]*(\d\.\d+)', line)
                if gpa_match:
                    edu['gpa'] = f"{gpa_match.group(1)}/{gpa_match.group(2)}"

                # Extract degree
                degree_match = re.search(r'(Bachelor|Master|PhD|학사|석사|박사|B\.S\.|M\.S\.)', line, re.I)
                if degree_match:
                    edu['degree'] = degree_match.group(1)

                # Extract dates
                dates = self._parse_date_range(line)
                edu.update(dates)

                # School name (simplified - remove dates/gpa/degree)
                school = re.sub(r'\d{4}[년\./-]?\s*\d{0,2}', '', line)
                school = re.sub(r'\d\.\d+/\d\.\d+', '', school)
                school = re.sub(r'(Bachelor|Master|PhD|학사|석사|박사)', '', school, flags=re.I)
                edu['school'] = school.strip()

                self.extracted['education'].append(edu)

    def _parse_skills(self, lines: List[str]):
        """Parse skills section"""
        skills = {'technical': [], 'languages': []}

        for line in lines:
            # Technical skills
            if any(kw in line.lower() for kw in ['python', 'java', 'framework', 'aws', 'pytorch', '개발']):
                # Clean bullet points
                clean_line = re.sub(r'^[•\-\*]\s*', '', line)
                skills['technical'].append({'value': clean_line, 'confidence': 0.8})

            # Languages
            if any(lang in line.lower() for lang in ['english', 'korean', 'japanese', 'chinese', '영어', '한국어']):
                skills['languages'].append({'value': line, 'confidence': 0.8})

        self.extracted['skills'] = skills

    def _process_table(self, table: Table):
        """Process table-based resume content"""
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            field_name = row.cells[0].text.strip()
            field_value = row.cells[1].text.strip()

            if not field_name or not field_value:
                continue

            # Map to personal info fields
            field_lower = field_name.lower()
            if any(kw in field_lower for kw in ['이름', '성명', 'name']):
                self.extracted['personal_info']['name'] = {
                    'value': field_value,
                    'confidence': 0.95,
                    'source': 'table'
                }
            elif any(kw in field_lower for kw in ['전화', '연락', 'phone', 'tel']):
                # Might contain phone and email together
                phone = re.search(r'(\d{2,3}[-\.\s]?\d{3,4}[-\.\s]?\d{4})', field_value)
                if phone:
                    self.extracted['personal_info']['phone'] = {
                        'value': phone.group(1),
                        'confidence': 0.9,
                        'source': 'table'
                    }
                email = re.search(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', field_value)
                if email:
                    self.extracted['personal_info']['email'] = {
                        'value': email.group(1),
                        'confidence': 0.9,
                        'source': 'table'
                    }
            elif 'email' in field_lower or '@' in field_value:
                email = re.search(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', field_value)
                if email:
                    self.extracted['personal_info']['email'] = {
                        'value': email.group(1),
                        'confidence': 0.9,
                        'source': 'table'
                    }
            elif any(kw in field_lower for kw in ['주소', 'address']):
                self.extracted['personal_info']['address'] = {
                    'value': field_value,
                    'confidence': 0.85,
                    'source': 'table'
                }

    def _extract_contact_info(self, text: str, source: str):
        """Extract contact information from text"""
        # Phone number
        phone = re.search(r'(\+?\d{1,3}[-\.\s]?)?\(?\d{2,3}\)?[-\.\s]?\d{3,4}[-\.\s]?\d{4}', text)
        if phone and 'phone' not in self.extracted['personal_info']:
            self.extracted['personal_info']['phone'] = {
                'value': phone.group(0),
                'confidence': 0.85,
                'source': source
            }

        # Email
        email = re.search(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', text)
        if email and 'email' not in self.extracted['personal_info']:
            self.extracted['personal_info']['email'] = {
                'value': email.group(1),
                'confidence': 0.9,
                'source': source
            }

        # Name (if at top and looks like a name)
        if source in ['header', 'body'] and 'name' not in self.extracted['personal_info']:
            # Simple heuristic: short text, 2-4 words, no numbers
            words = text.split()
            if 2 <= len(words) <= 4 and not any(char.isdigit() for char in text):
                if not any(kw in text.lower() for kw in ['phone', 'email', 'address']):
                    self.extracted['personal_info']['name'] = {
                        'value': text,
                        'confidence': 0.7,
                        'source': source
                    }

    def _deduplicate(self):
        """Remove duplicate information"""
        # If contact info found in multiple places, keep highest confidence
        for field in ['name', 'phone', 'email']:
            if field in self.extracted['personal_info']:
                if isinstance(self.extracted['personal_info'][field], list):
                    # Sort by confidence and keep best
                    sorted_items = sorted(
                        self.extracted['personal_info'][field],
                        key=lambda x: x.get('confidence', 0),
                        reverse=True
                    )
                    self.extracted['personal_info'][field] = sorted_items[0]

    def _calculate_total_experience_years(self) -> float:
        """Calculate total years of experience from all jobs"""
        try:
            from dateutil.relativedelta import relativedelta
            from datetime import datetime
        except ImportError:
            self.extracted['metadata']['warnings'].append(
                "dateutil not installed - cannot calculate total experience. Run: pip install python-dateutil"
            )
            return 0.0

        total_months = 0
        today = datetime.now()

        for job in self.extracted['experience']:
            start_str = job.get('start_date', '')
            end_str = job.get('end_date', 'Present')

            if not start_str:
                continue

            try:
                # Parse start date
                start_parts = start_str.split('.')
                start_date = datetime(int(start_parts[0]), int(start_parts[1]), 1)

                # Parse end date
                if end_str == 'Present':
                    end_date = today
                else:
                    end_parts = end_str.split('.')
                    end_date = datetime(int(end_parts[0]), int(end_parts[1]), 1)

                # Calculate months
                delta = relativedelta(end_date, start_date)
                months = delta.years * 12 + delta.months
                total_months += months

            except (ValueError, IndexError) as e:
                self.extracted['metadata']['warnings'].append(
                    f"Could not parse dates for experience: {start_str} - {end_str}"
                )

        return total_months / 12.0

    def _calculate_confidence(self):
        """Calculate overall extraction confidence"""
        scores = []

        # Personal info completeness
        required_personal = ['name', 'phone', 'email']
        personal_score = sum(1 for f in required_personal if f in self.extracted['personal_info']) / len(required_personal)
        scores.append(personal_score)

        # Experience found?
        scores.append(1.0 if self.extracted['experience'] else 0.5)

        # Education found?
        scores.append(1.0 if self.extracted['education'] else 0.7)

        # Skills found?
        scores.append(1.0 if self.extracted['skills'].get('technical') else 0.6)

        self.extracted['metadata']['overall_confidence'] = sum(scores) / len(scores)

        # Calculate total experience
        if self.extracted['experience']:
            total_years = self._calculate_total_experience_years()
            if total_years > 0:
                self.extracted['metadata']['total_experience_years'] = round(total_years, 1)

        # Warnings
        if not self.extracted['personal_info'].get('name'):
            self.extracted['metadata']['warnings'].append("Name not found")
        if not self.extracted['experience']:
            self.extracted['metadata']['warnings'].append("No work experience found")


def main():
    parser = argparse.ArgumentParser(description='Extract structured content from resumes')
    parser.add_argument('resume', help='Path to resume file (.docx)')
    parser.add_argument('--output', '-o', help='Output JSON file', default=None)
    args = parser.parse_args()

    # Parse resume
    resume_parser = ResumeParser(args.resume)
    result = resume_parser.parse()

    # Output
    output_json = json.dumps(result, indent=2, ensure_ascii=False)

    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(output_json)
        print(f"\nSaved to: {args.output}")
    else:
        print("\n" + output_json)

    # Summary
    print(f"\nExtraction Summary:")
    print(f"  Overall confidence: {result['metadata']['overall_confidence']:.1%}")
    print(f"  Personal info fields: {len(result['personal_info'])}")
    print(f"  Experience entries: {len(result['experience'])}")
    print(f"  Education entries: {len(result['education'])}")
    print(f"  Warnings: {len(result['metadata']['warnings'])}")
    if result['metadata']['warnings']:
        for warning in result['metadata']['warnings']:
            print(f"    - {warning}")


if __name__ == '__main__':
    main()
