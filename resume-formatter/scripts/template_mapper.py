#!/usr/bin/env python3
"""
Template Mapper - Map extracted content to target template (FIXED VERSION)

Maps extracted resume content to a target template while preserving quality.
Fixes critical bugs identified in evaluation:
1. LinkedIn duplication bug
2. Missing detailed experience mapping
3. Career summary not filled
4. Template placeholder cleanup

Usage:
    python template_mapper.py <parsed.json> <template.docx> [--output mapped.docx]
"""

import argparse
import json
import re
from pathlib import Path
from typing import Dict, Any, List, Set
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    exit(1)


class TemplateMapper:
    """Map extracted content to template with bug fixes"""

    # Field name mappings
    FIELD_MAPPINGS = {
        '성명': 'personal_info.name',
        'Name': 'personal_info.name',
        '이름': 'personal_info.name',
        '이    름': 'personal_info.name',
        '연락처': 'personal_info.contact',
        'Phone': 'personal_info.phone',
        '전화': 'personal_info.phone',
        '연 락 처': 'personal_info.phone',
        'Email': 'personal_info.email',
        '이 메 일': 'personal_info.email',
        '주소': 'personal_info.address',
        '주    소': 'personal_info.address',
        'Address': 'personal_info.address',
        '경력사항': 'experience',
        'Experience': 'experience',
        '학력사항': 'education',
        'Education': 'education',
        '핵심역량': 'skills.technical',
        'Skills': 'skills.technical',
    }

    # Template placeholder patterns to remove
    PLACEHOLDER_PATTERNS = [
        r'회사이름',
        r'부서명',
        r'직무',
        r'직급\(직책\)',
        r'ㅇㅇ대학교',
        r'ㅇㅇ제조',
        r'ㅇㅇㅇ',
        r'\d+,\d+억원',
        r'\d+,\d+만원',
        r'2000\.00',
        r'20XX\.',
        r'00년 00개월',
    ]

    def __init__(self, extracted_data: Dict[str, Any], template_path: str):
        self.data = extracted_data
        self.template_path = Path(template_path)
        self.doc = Document(str(self.template_path))
        self.added_urls = set()  # Track added URLs to prevent duplicates
        self.warnings = []
        self.mapped_sections = set()

    def map(self, output_path: str):
        """Map content to template and save"""
        print(f"Mapping content to template: {self.template_path.name}")

        # Step 1: Map personal information
        self._map_personal_info()

        # Step 2: Map education
        self._map_education()

        # Step 3: Map core competencies
        self._map_core_competencies()

        # Step 4: Map technical skills
        self._map_technical_skills()

        # Step 5: Map career summary (경력사항)
        self._map_career_summary()

        # Step 6: Map detailed experience (경력 세부사항) - CRITICAL FIX
        self._map_detailed_experience()

        # Step 7: Calculate and fill total experience
        self._fill_total_experience()

        # Step 8: Clean up template placeholders - CRITICAL FIX
        self._cleanup_template_placeholders()

        # Step 9: Handle optional sections
        self._handle_optional_sections()

        # Save
        self.doc.save(output_path)
        print(f"✓ Saved mapped resume to: {output_path}")

        # Print warnings
        if self.warnings:
            print(f"\n⚠ Warnings ({len(self.warnings)}):")
            for warning in self.warnings:
                print(f"  - {warning}")

    def _map_personal_info(self):
        """Map personal information fields"""
        print("  Mapping personal information...")

        for para in self.doc.paragraphs:
            text = para.text.strip()

            # Name
            if '이    름' in text or '성명' in text:
                name = self._get_value('personal_info.name')
                if name:
                    # Format: 이름 : 문다빈 (Dabin Moon) / 남 / 미혼
                    gender = self.data.get('personal_info', {}).get('gender', '정보 없음')
                    marital = self.data.get('personal_info', {}).get('marital_status', '정보 없음')
                    para.text = f"이    름 : {name} / {gender} / {marital}"

            # Birth date
            elif '생년월일' in text:
                birth = self._get_value('personal_info.birth_date')
                para.text = f"생년월일 : {birth if birth else '정보 없음'}"

            # Address
            elif '주    소' in text:
                address = self._get_value('personal_info.address')
                if address:
                    para.text = f"주    소 : {address}"

            # Phone
            elif '연 락 처' in text and '이 메 일' not in text:
                phone = self._get_value('personal_info.phone')
                if phone:
                    para.text = f"연 락 처 : {phone}"

            # Email
            elif '이 메 일' in text:
                email = self._get_value('personal_info.email')
                if email:
                    para.text = f"이 메 일 : {email}"

            # Salary info - mark as negotiable if not provided
            elif '현재연봉' in text:
                para.text = "현재연봉 : 협의 가능"
            elif '희망연봉' in text:
                para.text = "희망연봉 : 협의 가능"
            elif '입사시기' in text:
                para.text = "입사시기 : 협의 가능"

    def _map_education(self):
        """Map education section"""
        print("  Mapping education...")

        # Find education section
        for i, para in enumerate(self.doc.paragraphs):
            if '학력사항' in para.text:
                # Remove template placeholders (next few lines)
                self._clear_template_lines(i + 1, i + 5)

                # Insert actual education
                education = self.data.get('education', [])
                insert_idx = i + 1

                for edu in education:
                    # Format: 2014.09~2015.12   Columbia University 통계학 석사 졸업 (New York, NY, 3.4/4.0)
                    start = edu.get('start_date', '')
                    end = edu.get('end_date', '')
                    school = edu.get('school', '')
                    degree = edu.get('degree', '')
                    major = edu.get('major', '')
                    location = edu.get('location', '')
                    gpa = edu.get('gpa', '')

                    # Clean up school name
                    school = re.sub(r'\(.*?\)', '', school).strip()
                    school = re.sub(r'[-–].*$', '', school).strip()

                    date_range = f"{start}~{end}" if start and end else ''
                    gpa_str = f"({gpa})" if gpa else ''
                    location_str = f"({location})" if location else ''

                    edu_line = f"{date_range} \t{school} {major} {degree} 졸업 {location_str} {gpa_str}".strip()

                    # Insert as new paragraph
                    new_para = self.doc.paragraphs[insert_idx].insert_paragraph_before(edu_line)
                    insert_idx += 1

                self.mapped_sections.add('education')
                break

    def _map_core_competencies(self):
        """Map core competencies (핵심역량)"""
        print("  Mapping core competencies...")

        for i, para in enumerate(self.doc.paragraphs):
            if '핵심역량' in para.text and i < len(self.doc.paragraphs) - 1:
                # Remove template examples
                self._clear_template_lines(i + 1, i + 5)

                # Generate competency summary from experience
                competencies = self._generate_competencies()

                # Insert competencies
                insert_idx = i + 1
                for comp in competencies:
                    new_para = self.doc.paragraphs[insert_idx].insert_paragraph_before(comp)
                    insert_idx += 1

                self.mapped_sections.add('core_competencies')
                break

    def _generate_competencies(self) -> List[str]:
        """Generate core competencies from experience and skills"""
        competencies = []

        # Get experience summary
        experience = self.data.get('experience', [])
        if experience:
            # Calculate total years
            total_years = self._calculate_years_experience()

            # Get domain from most recent role
            recent = experience[0] if experience else {}
            domain = self._infer_domain(recent)

            # Competency 1: Domain expertise
            comp1 = f"{domain} 분야에서 {total_years}년 이상 연구/개발 경험"
            if recent.get('responsibilities'):
                # Add key achievements
                key_projects = self._extract_key_projects(recent)
                if key_projects:
                    comp1 += f" ({', '.join(key_projects[:3])})"
            competencies.append(comp1)

        # Competency 2: Technical stack
        skills = self.data.get('skills', {}).get('technical', [])
        if skills:
            # Get ML/DL frameworks
            ml_tools = [s.get('value', s) if isinstance(s, dict) else s for s in skills
                       if any(kw in str(s).lower() for kw in ['pytorch', 'tensorflow', 'ml', 'dl'])]
            cloud = [s.get('value', s) if isinstance(s, dict) else s for s in skills
                    if any(kw in str(s).lower() for kw in ['aws', 'gcp', 'cloud'])]

            if ml_tools or cloud:
                comp2 = f"{', '.join(ml_tools[:2])} 기반 ML/DL 모델 개발"
                if cloud:
                    comp2 += f"부터 {', '.join(cloud[:2])} 기반 MLOps 파이프라인 구축"
                competencies.append(comp2)

        # Competency 3: Product/service experience
        if experience:
            products = []
            for exp in experience:
                # Extract product names from responsibilities
                for resp in exp.get('responsibilities', []):
                    # Look for product names (usually in English or with URLs)
                    product_match = re.findall(r'([A-Z][a-z]+(?:[A-Z][a-z]+)*)', resp)
                    products.extend([p for p in product_match if len(p) > 3])

            if products:
                unique_products = list(dict.fromkeys(products))[:3]  # Deduplicate
                comp3 = f"{', '.join(unique_products)} 개발 리딩 및 상용화 경험"
                competencies.append(comp3)

        return competencies if competencies else ["정보 없음"]

    def _map_technical_skills(self):
        """Map technical skills (기술스택)"""
        print("  Mapping technical skills...")

        for i, para in enumerate(self.doc.paragraphs):
            if '기술스택' in para.text:
                # Remove template examples
                self._clear_template_lines(i + 1, i + 10)

                # Get skills
                skills = self.data.get('skills', {}).get('technical', [])

                # Categorize skills
                categorized = self._categorize_skills(skills)

                # Insert categorized skills
                insert_idx = i + 1
                for category, items in categorized.items():
                    if items:
                        skill_line = f"{category}: {', '.join(items[:10])}"  # Limit to 10 per category
                        new_para = self.doc.paragraphs[insert_idx].insert_paragraph_before(skill_line)
                        insert_idx += 1

                # Add portfolio/LinkedIn - FIX: Add only once
                linkedin = self._get_value('personal_info.linkedin')
                github = self._get_value('personal_info.github')
                portfolio = self._get_value('personal_info.portfolio')

                if linkedin and linkedin not in self.added_urls:
                    self.doc.paragraphs[insert_idx].insert_paragraph_before(f"LinkedIn: {linkedin}")
                    self.added_urls.add(linkedin)
                    insert_idx += 1

                if github and github not in self.added_urls:
                    self.doc.paragraphs[insert_idx].insert_paragraph_before(f"GitHub: {github}")
                    self.added_urls.add(github)
                    insert_idx += 1

                if portfolio and portfolio not in self.added_urls:
                    self.doc.paragraphs[insert_idx].insert_paragraph_before(f"Portfolio: {portfolio}")
                    self.added_urls.add(portfolio)
                    insert_idx += 1

                self.mapped_sections.add('technical_skills')
                break

    def _categorize_skills(self, skills: List) -> Dict[str, List[str]]:
        """Categorize skills into groups"""
        categories = {
            'Python, R, SQL/NoSQL': [],
            'Cloud': [],
            'ML/DL': [],
            'MLOps': [],
            'LLM/Agents': [],
            'Audio Processing': [],
            'Data': []
        }

        for skill in skills:
            value = skill.get('value', skill) if isinstance(skill, dict) else str(skill)
            value_lower = value.lower()

            if any(kw in value_lower for kw in ['python', 'sql']):
                categories['Python, R, SQL/NoSQL'].append(value)
            elif any(kw in value_lower for kw in ['aws', 'gcp', 'cloud', 'azure']):
                categories['Cloud'].append(value)
            elif any(kw in value_lower for kw in ['pytorch', 'tensorflow', 'sklearn', 'scikit']):
                categories['ML/DL'].append(value)
            elif any(kw in value_lower for kw in ['git', 'docker', 'kubernetes', 'mlops']):
                categories['MLOps'].append(value)
            elif any(kw in value_lower for kw in ['llm', 'langchain', 'claude', 'gpt', 'rag']):
                categories['LLM/Agents'].append(value)
            elif any(kw in value_lower for kw in ['audio', 'librosa', 'torchaudio']):
                categories['Audio Processing'].append(value)
            elif any(kw in value_lower for kw in ['spark', 'hadoop', 'data']):
                categories['Data'].append(value)

        # Remove empty categories
        return {k: v for k, v in categories.items() if v}

    def _map_career_summary(self):
        """Map career summary (경력사항) - CRITICAL FIX"""
        print("  Mapping career summary...")

        for i, para in enumerate(self.doc.paragraphs):
            if '경력사항' in para.text and '세부사항' not in para.text:
                # Update total experience in header
                total_exp = self._calculate_total_experience_formatted()
                if '총 경력' in para.text:
                    para.text = para.text.replace('00년 00개월', total_exp)

                # Remove template placeholder lines
                self._clear_template_lines(i + 1, i + 5)

                # Insert actual career entries
                experience = self.data.get('experience', [])
                insert_idx = i + 1

                for exp in experience:
                    start = exp.get('start_date', '')
                    end = exp.get('end_date', '재직 중')
                    company = exp.get('company', '회사명')
                    title = exp.get('title', '직무')

                    # Format: 2023.07~재직 중	Hudson AI / AI Research / AI Researcher
                    career_line = f"{start}~{end}\t{company} / {title}"

                    new_para = self.doc.paragraphs[insert_idx].insert_paragraph_before(career_line)
                    insert_idx += 1

                self.mapped_sections.add('career_summary')
                print(f"    ✓ Added {len(experience)} career entries")
                break

    def _map_detailed_experience(self):
        """Map detailed experience section (경력 세부사항) - CRITICAL FIX"""
        print("  Mapping detailed experience...")

        for i, para in enumerate(self.doc.paragraphs):
            if '경력 세부사항' in para.text or '세부사항' in para.text:
                # Remove ALL template placeholder content
                self._clear_template_lines(i + 1, i + 40)

                # Insert actual detailed experience
                experience = self.data.get('experience', [])
                insert_idx = i + 1

                for exp_idx, exp in enumerate(experience):
                    # Add job block
                    job_lines = self._format_detailed_job(exp, exp_idx)

                    for line in job_lines:
                        new_para = self.doc.paragraphs[insert_idx].insert_paragraph_before(line)
                        insert_idx += 1

                    # Add spacing between jobs
                    self.doc.paragraphs[insert_idx].insert_paragraph_before("")
                    insert_idx += 1

                self.mapped_sections.add('detailed_experience')
                print(f"    ✓ Added detailed experience for {len(experience)} roles")
                return

        # If section not found, warn
        self.warnings.append("경력 세부사항 section not found in template")

    def _format_detailed_job(self, exp: Dict[str, Any], index: int) -> List[str]:
        """Format a detailed job entry"""
        lines = []

        # Header: dates and company/title
        start = exp.get('start_date', '2000.00')
        end = exp.get('end_date', '재직 중')
        company = exp.get('company', '회사명')
        title = exp.get('title', '직무')

        lines.append(f"{start}~ {end}    {company} / {title}")

        # Company description (if we can infer it)
        if company and company != '회사명':
            lines.append("[회사소개]")
            # Try to infer industry from responsibilities
            industry = self._infer_industry(exp)
            lines.append(f"- 업종 및 제품 : {industry}")

        # Main responsibilities
        lines.append("[주요 업무/성과]")

        responsibilities = exp.get('responsibilities', [])
        if responsibilities:
            # Group by projects if possible
            current_project = None
            project_num = 1

            for resp in responsibilities[:15]:  # Limit to top 15 items
                # Check if this looks like a project header
                if self._looks_like_project(resp):
                    lines.append(f"{project_num}. {resp}")
                    current_project = resp
                    project_num += 1
                else:
                    # Regular responsibility - add as bullet
                    lines.append(f"- {resp}")
        else:
            lines.append("- 정보 없음")

        # Optional: reason for leaving (if not current job)
        if end != '재직 중' and end != 'Present':
            lines.append("이직사유 : 협의 가능")

        return lines

    def _looks_like_project(self, text: str) -> bool:
        """Check if text looks like a project name/header"""
        # Project indicators
        indicators = ['구축', '연구', '개발', '프로젝트', '서비스', 'PoC']
        # Has dates in parentheses
        has_dates = re.search(r'\(20\d{2}\.\d{2}\s*-\s*.*?\)', text)

        return has_dates or any(ind in text for ind in indicators) and len(text) < 100

    def _fill_total_experience(self):
        """Fill total experience duration"""
        total = self._calculate_total_experience_formatted()

        for para in self.doc.paragraphs:
            if '총 경력' in para.text and '00년 00개월' in para.text:
                para.text = para.text.replace('00년 00개월', total)

    def _calculate_total_experience_formatted(self) -> str:
        """Calculate total experience in years and months"""
        experience = self.data.get('experience', [])
        if not experience:
            return "00년 00개월"

        from dateutil.relativedelta import relativedelta

        total_months = 0
        for exp in experience:
            start_str = exp.get('start_date', '')
            end_str = exp.get('end_date', 'Present')

            try:
                start = self._parse_date_str(start_str)
                end = datetime.now() if end_str in ['Present', '현재', '재직중'] else self._parse_date_str(end_str)

                if start and end:
                    delta = relativedelta(end, start)
                    total_months += delta.years * 12 + delta.months
            except:
                continue

        years = total_months // 12
        months = total_months % 12
        return f"{years:02d}년 {months:02d}개월"

    def _parse_date_str(self, date_str: str):
        """Parse date string to datetime"""
        if not date_str:
            return None

        # Try YYYY.MM
        match = re.match(r'(\d{4})\.(\d{2})', date_str)
        if match:
            return datetime(int(match.group(1)), int(match.group(2)), 1)

        # Try YYYY-MM
        match = re.match(r'(\d{4})-(\d{2})', date_str)
        if match:
            return datetime(int(match.group(1)), int(match.group(2)), 1)

        return None

    def _cleanup_template_placeholders(self):
        """Remove template placeholder text - CRITICAL FIX"""
        print("  Cleaning up template placeholders...")

        removed_count = 0
        for para in self.doc.paragraphs:
            text = para.text

            # Check if contains placeholder patterns
            is_placeholder = False
            for pattern in self.PLACEHOLDER_PATTERNS:
                if re.search(pattern, text):
                    is_placeholder = True
                    break

            if is_placeholder:
                # Check if this is a section we didn't map
                if not any(section in text for section in ['개인정보', '학력사항', '경력사항', '핵심역량', '기술스택']):
                    para.text = ""
                    removed_count += 1

        print(f"    ✓ Removed {removed_count} placeholder lines")

    def _handle_optional_sections(self):
        """Handle optional sections (병역, 어학, 자격증, etc.)"""
        print("  Handling optional sections...")

        for para in self.doc.paragraphs:
            text = para.text

            # Military service
            if '병  역' in text and '2000.00' in text:
                military = self.data.get('personal_info', {}).get('military_service')
                if military:
                    para.text = f"병  역 : {military}"
                else:
                    para.text = "병  역 : 해당 없음"

            # Language proficiency
            elif '어  학' in text and ('Business' in text or '2000.00' in text):
                languages = self.data.get('skills', {}).get('languages', [])
                if languages:
                    lang_str = ', '.join([l.get('value', l) if isinstance(l, dict) else str(l) for l in languages[:2]])
                    para.text = f"어  학 : {lang_str}"
                else:
                    para.text = ""  # Remove if no data

            # Certifications
            elif '자격증' in text and ('빅데이터' in text or '2021' in text):
                certs = self.data.get('certifications', [])
                if certs:
                    cert_str = ', '.join([c.get('name', c) if isinstance(c, dict) else str(c) for c in certs[:3]])
                    para.text = f"자격증 : {cert_str}"
                else:
                    para.text = ""  # Remove if no data

            # Self-introduction template
            elif '[자유양식]' in text or '지원동기와 입사 후 포부' in text:
                para.text = ""  # Clear template guidance

    def _clear_template_lines(self, start_idx: int, end_idx: int):
        """Clear template placeholder lines"""
        for i in range(start_idx, min(end_idx, len(self.doc.paragraphs))):
            if i < len(self.doc.paragraphs):
                self.doc.paragraphs[i].text = ""

    def _get_value(self, field_path: str) -> Any:
        """Get value from extracted data by path"""
        parts = field_path.split('.')
        value = self.data

        for part in parts:
            if isinstance(value, dict):
                value = value.get(part)
            else:
                return None

            if value is None:
                return None

        # Handle different value types
        if isinstance(value, dict) and 'value' in value:
            return value['value']
        elif isinstance(value, list) and value:
            return value
        else:
            return value

    def _calculate_years_experience(self) -> int:
        """Calculate total years of experience"""
        total = self._calculate_total_experience_formatted()
        match = re.match(r'(\d+)년', total)
        return int(match.group(1)) if match else 0

    def _infer_domain(self, exp: Dict) -> str:
        """Infer domain from experience"""
        responsibilities = ' '.join(exp.get('responsibilities', []))

        if any(kw in responsibilities for kw in ['TTS', '음성', 'Audio', 'Speech']):
            return "음성합성 및 대화형 AI"
        elif any(kw in responsibilities for kw in ['LLM', 'GPT', 'Claude', '대화']):
            return "LLM 및 대화형 AI"
        elif any(kw in responsibilities for kw in ['ML', 'DL', '머신러닝', '딥러닝']):
            return "ML/DL"
        else:
            return "AI/ML"

    def _infer_industry(self, exp: Dict) -> str:
        """Infer industry from experience"""
        company = exp.get('company', '')
        responsibilities = ' '.join(exp.get('responsibilities', []))

        if 'AI' in company or 'AI' in responsibilities:
            return "AI/ML 기술 개발 및 서비스"
        elif any(kw in responsibilities for kw in ['게임', 'Game']):
            return "게임 개발 및 운영"
        else:
            return "소프트웨어 개발 및 서비스"

    def _extract_key_projects(self, exp: Dict) -> List[str]:
        """Extract key project names from responsibilities"""
        projects = []
        for resp in exp.get('responsibilities', []):
            # Look for capitalized product names
            matches = re.findall(r'\b([A-Z][a-z]+(?:[A-Z][a-z]+)*)\b', resp)
            projects.extend([m for m in matches if len(m) > 3 and m not in ['Text', 'Speech', 'Auto']])

        return list(dict.fromkeys(projects))  # Deduplicate while preserving order


def main():
    parser = argparse.ArgumentParser(description='Map extracted content to template (FIXED)')
    parser.add_argument('extracted_json', help='Path to parsed resume JSON')
    parser.add_argument('template', help='Path to template file (.docx)')
    parser.add_argument('--output', '-o', help='Output file', required=True)
    args = parser.parse_args()

    # Load extracted data
    with open(args.extracted_json, 'r', encoding='utf-8') as f:
        extracted = json.load(f)

    # Map to template
    mapper = TemplateMapper(extracted, args.template)
    mapper.map(args.output)

    print(f"\n✓ Mapping complete!")
    print(f"  Mapped sections: {', '.join(mapper.mapped_sections)}")


if __name__ == '__main__':
    main()
