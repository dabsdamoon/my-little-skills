#!/usr/bin/env python3
"""
Content Validator - Validate resume transformation quality

Compares source and target to ensure completeness and quality.

Usage:
    python content_validator.py <source_parsed.json> <target_resume.docx>
"""

import argparse
import json
import re
from pathlib import Path
from typing import Dict, Any, List

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    exit(1)


class ContentValidator:
    """Validate resume transformation"""

    def __init__(self, source_data: Dict[str, Any], target_path: str):
        self.source = source_data
        self.target_path = Path(target_path)
        self.target_doc = Document(str(self.target_path))
        self.target_text = self._extract_target_text()
        self.issues = []

    def validate(self) -> Dict[str, Any]:
        """Run all validation checks"""
        print(f"Validating: {self.target_path.name}")

        report = {
            'overall_score': 0.0,
            'completeness': {},
            'quality_checks': {},
            'issues': [],
            'recommendations': []
        }

        # Check completeness
        report['completeness'] = self._check_completeness()

        # Check quality
        self._check_quality()

        # NEW: Check for duplicates
        report['quality_checks']['duplicates'] = self._check_for_duplicates()

        # NEW: Check required sections
        report['quality_checks']['required_sections'] = self._check_required_sections()

        # NEW: Check template placeholders
        report['quality_checks']['template_cleaned'] = self._check_template_placeholders()

        # NEW: Check total experience years match
        report['quality_checks']['experience_years_match'] = self._check_experience_years()

        # Calculate overall score (weight completeness more)
        completeness_avg = sum(report['completeness'].values()) / len(report['completeness']) if report['completeness'] else 0.0
        quality_avg = sum(report['quality_checks'].values()) / len(report['quality_checks']) if report['quality_checks'] else 1.0
        report['overall_score'] = (completeness_avg * 0.7) + (quality_avg * 0.3)

        # Add issues and recommendations
        report['issues'] = self.issues

        # Generate recommendations
        report['recommendations'] = self._generate_recommendations(report)

        return report

    def _extract_target_text(self) -> str:
        """Extract all text from target document"""
        paragraphs = [p.text for p in self.target_doc.paragraphs]
        tables = []
        for table in self.target_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    tables.append(cell.text)

        return '\n'.join(paragraphs + tables)

    def _check_completeness(self) -> Dict[str, float]:
        """Check if all source content is in target"""
        completeness = {}

        # Personal info
        personal_fields = ['name', 'phone', 'email']
        personal_found = 0
        for field in personal_fields:
            value = self.source.get('personal_info', {}).get(field, {})
            if isinstance(value, dict):
                value = value.get('value', '')
            if value and value in self.target_text:
                personal_found += 1
            elif value:
                self.issues.append(f"Personal info '{field}' not found in target: {value}")

        completeness['personal_info'] = personal_found / len(personal_fields) if personal_fields else 0.0

        # Experience
        exp_count = len(self.source.get('experience', []))
        exp_found = 0
        for exp in self.source.get('experience', []):
            company = exp.get('company', '')
            if company and company in self.target_text:
                exp_found += 1
            elif company:
                self.issues.append(f"Experience company not found: {company}")

        completeness['experience'] = exp_found / exp_count if exp_count > 0 else 1.0

        # Education
        edu_count = len(self.source.get('education', []))
        edu_found = 0
        for edu in self.source.get('education', []):
            school = edu.get('school', '')
            if school and school in self.target_text:
                edu_found += 1

        completeness['education'] = edu_found / edu_count if edu_count > 0 else 0.8

        # Skills
        skills = self.source.get('skills', {}).get('technical', [])
        skill_found = 0
        for skill in skills[:10]:  # Check top 10
            value = skill.get('value', skill) if isinstance(skill, dict) else skill
            if value and value in self.target_text:
                skill_found += 1

        completeness['skills'] = skill_found / min(len(skills), 10) if skills else 0.8

        return completeness

    def _check_quality(self):
        """Check content quality"""
        # Check for encoding issues (garbled Korean text)
        if '�' in self.target_text:
            self.issues.append("Encoding issues detected (garbled characters)")

    def _check_for_duplicates(self) -> float:
        """Check for duplicate content (e.g., repeated URLs, company names)"""
        import re
        from collections import Counter

        # Extract URLs
        urls = re.findall(r'https?://[^\s]+', self.target_text)
        url_counts = Counter(urls)
        duplicates_found = False

        for url, count in url_counts.items():
            if count > 1:
                self.issues.append(f"CRITICAL: URL duplicated {count} times: {url}")
                duplicates_found = True

        # Check for repeated long phrases (10+ chars repeated 3+ times)
        lines = [line.strip() for line in self.target_text.split('\n') if len(line.strip()) > 10]
        line_counts = Counter(lines)

        for line, count in line_counts.items():
            if count >= 3 and line not in ['', ' ']:
                # Ignore common headers
                if not any(header in line for header in ['인적사항', '학력사항', '경력사항']):
                    self.issues.append(f"Repeated content ({count}x): {line[:50]}...")
                    duplicates_found = True

        return 0.0 if duplicates_found else 1.0

    def _check_required_sections(self) -> float:
        """Check if required sections are filled (not just template text)"""
        required_sections = {
            '경력 세부사항': ['회사소개', '주요 업무', '담당 업무'],
            '경력사항': ['회사', '직무', '기간'],
            '인적사항': ['이름', '연락처', 'email'],
        }

        score = 1.0
        for section_name, required_content_types in required_sections.items():
            # Find section in document
            section_found = False
            section_has_content = False

            for para_idx, para in enumerate(self.target_doc.paragraphs):
                if section_name in para.text:
                    section_found = True
                    # Check next 20 paragraphs for actual content
                    section_text = '\n'.join([
                        p.text for p in self.target_doc.paragraphs[para_idx:para_idx+20]
                    ])

                    # Check if section has real content (not just template)
                    has_real_content = False
                    for content_type in required_content_types:
                        # Look for actual values, not just the label
                        if content_type in section_text:
                            # Check if there's text after the label
                            pattern = f"{content_type}[:\\s]+(\\S.+)"
                            match = re.search(pattern, section_text)
                            if match and len(match.group(1).strip()) > 3:
                                has_real_content = True
                                break

                    if has_real_content:
                        section_has_content = True
                    else:
                        self.issues.append(f"CRITICAL: Section '{section_name}' is empty or has only template text")
                        score -= 0.3

                    break

            if not section_found:
                self.issues.append(f"WARNING: Required section '{section_name}' not found")
                score -= 0.2

        return max(0.0, score)

    def _check_template_placeholders(self) -> float:
        """Check if template placeholder text has been removed"""
        # Common Korean template placeholders
        placeholders = [
            'OOO', 'XXX', '[Insert', '정보 없음', 'N/A',
            '여기에', '입력하세요', '작성하세요',
            '예시:', '샘플', 'sample',
            '회사이름 / 부서명 / 직무',  # Specific from the bug
            '주식회사 ABC', '주식회사 DEF',  # Fake companies
        ]

        found_placeholders = []
        for placeholder in placeholders:
            if placeholder in self.target_text:
                found_placeholders.append(placeholder)
                self.issues.append(f"Template placeholder not removed: '{placeholder}'")

        return 0.0 if found_placeholders else 1.0

    def _check_experience_years(self) -> float:
        """Check if total experience years match between source and target"""
        import re
        from dateutil.relativedelta import relativedelta
        from datetime import datetime

        # Get source total years
        source_total = self.source.get('metadata', {}).get('total_experience_years', 0)
        if source_total == 0:
            self.issues.append("WARNING: Source total experience years not calculated")
            return 0.8

        # Extract companies and dates from target document
        target_companies = []
        target_total_months = 0

        # Look for career summary section (경력사항)
        for i, para in enumerate(self.target_doc.paragraphs):
            if '경력사항' in para.text or '경력 사항' in para.text:
                # Parse next 20 lines for company/date entries
                for j in range(i+1, min(i+20, len(self.target_doc.paragraphs))):
                    text = self.target_doc.paragraphs[j].text.strip()

                    # Look for date range patterns
                    # Format: 2023.07~Present	Hudson AI / AI Researcher/Engineer
                    match = re.search(r'(\d{4}\.\d{2})\s*[~-]\s*(Present|현재|\d{4}\.\d{2})\s+(.+)', text)
                    if match:
                        start_str = match.group(1)
                        end_str = match.group(2)
                        company_info = match.group(3)

                        # Extract company name (before / or ,)
                        company_name = company_info.split('/')[0].split(',')[0].strip()
                        target_companies.append(company_name)

                        # Calculate months
                        try:
                            start_parts = start_str.split('.')
                            start_date = datetime(int(start_parts[0]), int(start_parts[1]), 1)

                            if end_str in ['Present', '현재']:
                                end_date = datetime.now()
                            else:
                                end_parts = end_str.split('.')
                                end_date = datetime(int(end_parts[0]), int(end_parts[1]), 1)

                            delta = relativedelta(end_date, start_date)
                            months = delta.years * 12 + delta.months
                            target_total_months += months
                        except (ValueError, IndexError):
                            pass

                break

        target_total_years = target_total_months / 12.0 if target_total_months > 0 else 0

        # Get source companies
        source_companies = [exp.get('company', '') for exp in self.source.get('experience', [])]

        # Compare
        score = 1.0

        if target_total_years == 0:
            self.issues.append("CRITICAL: Could not calculate total experience from target document")
            score = 0.0
        else:
            # Check if years are close (within 20% tolerance)
            diff = abs(source_total - target_total_years)
            tolerance = source_total * 0.2  # 20% tolerance

            if diff > tolerance:
                self.issues.append(
                    f"CRITICAL: Experience years mismatch - Source: {source_total:.1f} years, "
                    f"Target: {target_total_years:.1f} years (diff: {diff:.1f} years)"
                )
                score -= 0.5

            # Check for missing companies
            missing_companies = []
            for src_company in source_companies:
                # Check if any target company contains source company name
                found = False
                for tgt_company in target_companies:
                    if src_company.lower() in tgt_company.lower() or tgt_company.lower() in src_company.lower():
                        found = True
                        break
                if not found:
                    missing_companies.append(src_company)

            if missing_companies:
                self.issues.append(
                    f"CRITICAL: {len(missing_companies)} companies missing from target: {', '.join(missing_companies)}"
                )
                score -= (0.2 * len(missing_companies))

        return max(0.0, score)

    def _generate_recommendations(self, report: Dict[str, Any]) -> List[str]:
        """Generate actionable recommendations based on validation results"""
        recommendations = []

        # Check completeness
        if report['overall_score'] < 0.9:
            recommendations.append("Review flagged missing content above")

        # Check for critical issues
        critical_issues = [i for i in report['issues'] if 'CRITICAL' in i]
        if critical_issues:
            recommendations.append(f"Fix {len(critical_issues)} critical issues immediately")

        # Check experience years
        if report['quality_checks'].get('experience_years_match', 1.0) < 0.8:
            recommendations.append("Review missing companies and experience year mismatch")

        # Check duplicates
        if report['quality_checks'].get('duplicates', 1.0) < 1.0:
            recommendations.append("Remove duplicate URLs and repeated content")

        # Check sections
        if report['quality_checks'].get('required_sections', 1.0) < 0.8:
            recommendations.append("Fill in empty required sections (경력 세부사항, 경력사항)")

        # Check placeholders
        if report['quality_checks'].get('template_cleaned', 1.0) < 1.0:
            recommendations.append("Remove all template placeholder text")

        # General recommendations
        if not recommendations:
            recommendations.append("Verify critical fields (name, contact) are correct")
            recommendations.append("Check for any truncated text")

        return recommendations

    def print_report(self, report: Dict[str, Any]):
        """Print validation report"""
        print("\n" + "="*60)
        print("VALIDATION REPORT")
        print("="*60)

        print(f"\nOverall Score: {report['overall_score']:.1%}")

        print("\nSection Completeness:")
        for section, score in report['completeness'].items():
            status = "✓" if score >= 0.9 else "⚠" if score >= 0.7 else "✗"
            print(f"  {status} {section}: {score:.1%}")

        print("\nQuality Checks:")
        for check, score in report.get('quality_checks', {}).items():
            status = "✓" if score >= 0.9 else "✗"
            print(f"  {status} {check}: {score:.1%}")

        if report['issues']:
            print(f"\nIssues Found ({len(report['issues'])}):")
            for issue in report['issues']:
                print(f"  ⚠ {issue}")
        else:
            print("\n✓ No issues found!")

        # Recommendations
        if report.get('recommendations'):
            print("\nRecommendations:")
            for rec in report['recommendations']:
                print(f"  → {rec}")


def main():
    parser = argparse.ArgumentParser(description='Validate resume transformation')
    parser.add_argument('source_json', help='Path to source parsed JSON')
    parser.add_argument('target_resume', help='Path to target resume (.docx)')
    parser.add_argument('--output', '-o', help='Output JSON report', default=None)
    args = parser.parse_args()

    # Load source data
    with open(args.source_json, 'r', encoding='utf-8') as f:
        source = json.load(f)

    # Validate
    validator = ContentValidator(source, args.target_resume)
    report = validator.validate()

    # Print report
    validator.print_report(report)

    # Save if requested
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
        print(f"\nReport saved to: {args.output}")


if __name__ == '__main__':
    main()
